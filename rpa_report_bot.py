"""
rpa_report_bot.py: the RPA bot (Solution 1 of the IAPA CA).

Automates the repetitive, rule-based "Daily Customer Creation" report
(report #4 of the Glo Customer Care Report List): read customers.csv ->
validate every record against the field-catalogue rules -> keep the ones
created on the report date -> reconcile the counts -> write the formatted
report to .xlsx AND .docx. Rejected records go to an exceptions csv for a
human, the RPA equivalent of a boundary error event.

Run:  python3 rpa_report_bot.py                       (report date = latest in the file)
      python3 rpa_report_bot.py --date 2026-07-06     (pick a specific day)
"""

import argparse
from datetime import datetime

import pandas as pd
from docx import Document

from Config import (
    CUSTOMERS_PATH,
    REQUIRED_CUSTOMER_COLUMNS,
    MSISDN_LENGTH,
    VALID_PREFIXES,
    VALID_SEGMENTS,
    VALID_STATUSES,
    REPORT_XLSX_PATH,
    REPORT_DOCX_PATH,
    EXCEPTIONS_PATH,
)


def load_customers(input_path):
    print(f"loading customer records from {input_path}...")

    #read the customers csv, keep MSISDN as text so leading zeros survive
    df = pd.read_csv(input_path, dtype={'MSISDN': str})

    #validation, make sure the columns we need are actually there
    for col in REQUIRED_CUSTOMER_COLUMNS:
        if col not in df.columns:
            raise ValueError(f"Required column '{col}' is missing from the customers csv.")

    print(f"loaded {len(df)} records.")
    return df


def validate_records(df):
    print("\nvalidating every record against the field-catalogue rules...")

    #collect the broken rules for each record, one string per record
    problems = []
    for _, row in df.iterrows():
        record_problems = []

        #rule 1: no required field may be empty
        for col in REQUIRED_CUSTOMER_COLUMNS:
            if pd.isna(row[col]) or str(row[col]).strip() == '':
                record_problems.append(f"{col} is empty")

        #rule 2: MSISDN must be 11 digits and start with a Glo prefix (0805, 0705..., not 0999)
        msisdn = str(row['MSISDN'])
        if len(msisdn) != MSISDN_LENGTH or not msisdn.isdigit():
            record_problems.append(f"MSISDN '{msisdn}' is not {MSISDN_LENGTH} digits")
        elif msisdn[:4] not in VALID_PREFIXES:
            record_problems.append(f"MSISDN prefix '{msisdn[:4]}' not in catalogue")

        #rule 3: segment must be one the catalogue knows
        if row['CustomerSegment'] not in VALID_SEGMENTS:
            record_problems.append(f"unknown segment '{row['CustomerSegment']}'")

        #rule 4: status must be one the catalogue knows
        if row['Status'] not in VALID_STATUSES:
            record_problems.append(f"unknown status '{row['Status']}'")

        #rule 5: creation date must be a real YYYY-MM-DD date, not garbage like "07/07/2026"
        try:
            datetime.strptime(str(row['CreationDate']), '%Y-%m-%d')
        except ValueError:
            record_problems.append(f"bad CreationDate '{row['CreationDate']}'")

        #join this record's broken rules into one cell ('' = clean record)
        problems.append('; '.join(record_problems))

    df['validation_errors'] = problems

    #the flow; clean rows go on to the report, broken rows leave through the side door
    df_valid = df[df['validation_errors'] == ''].drop(columns='validation_errors')
    df_exceptions = df[df['validation_errors'] != '']

    print(f"validation complete: {len(df_valid)} clean, {len(df_exceptions)} exceptions.")
    return df_valid, df_exceptions


def filter_report_day(df_valid, report_date):
    #default report date = the latest creation date in the file
    if report_date is None:
        report_date = df_valid['CreationDate'].max()

    #keep only the customers created on the report date
    df_day = df_valid[df_valid['CreationDate'] == report_date]

    print(f"\nreport date {report_date}: {len(df_day)} new customers.")
    return df_day, report_date


def summarise(df_day):
    #reconcile the counts, totals by segment and by status
    by_segment = df_day['CustomerSegment'].value_counts().rename_axis('CustomerSegment').reset_index(name='NewCustomers')
    by_status = df_day['Status'].value_counts().rename_axis('Status').reset_index(name='NewCustomers')

    print("count by segment:")
    print(by_segment.to_string(index=False))

    return by_segment, by_status


def write_excel(df_day, by_segment, by_status, report_date, output_path):
    print(f"\nwriting excel report to {output_path}...")

    #one workbook, three sheets: the detail rows + the two reconciled summaries
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        df_day.to_excel(writer, sheet_name='New Customers', index=False)
        by_segment.to_excel(writer, sheet_name='Summary by Segment', index=False)
        by_status.to_excel(writer, sheet_name='Summary by Status', index=False)

    print("✓ excel report saved")


def write_docx(df_day, by_segment, report_date, output_path):
    print(f"writing word report to {output_path}...")

    doc = Document()

    #title + the reconciliation line the ops team reads first
    doc.add_heading('Daily Customer Creation Report', level=1)
    doc.add_paragraph(f"Report date: {report_date}")
    doc.add_paragraph(f"Total new customers created: {len(df_day)}")
    doc.add_paragraph(f"Generated automatically on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} by rpa_report_bot v1.")

    #summary table: one header row + one row per segment
    doc.add_heading('New customers by segment', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Customer Segment'
    table.rows[0].cells[1].text = 'New Customers'
    for _, row in by_segment.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['CustomerSegment'])
        cells[1].text = str(row['NewCustomers'])

    doc.save(output_path)
    print("✓ word report saved")


def save_exceptions(df_exceptions, output_path):
    #the paper trail; rejected records go to a human, nothing gets dropped silently
    df_exceptions.to_csv(output_path, index=False)
    print(f"✓ {len(df_exceptions)} exception records saved to {output_path}")


#MAIN EXECUTION
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate the Daily Customer Creation report from customers.csv.")
    parser.add_argument('--input', type=str, default=CUSTOMERS_PATH, help="Path to input customers csv")
    parser.add_argument('--date', type=str, default=None, help="Report date YYYY-MM-DD (default: latest in file)")
    args = parser.parse_args()

    #1. load the raw customer records
    df = load_customers(args.input)

    #2. validate every record against the catalogue rules
    df_valid, df_exceptions = validate_records(df)

    #3. keep only the report day's new customers
    df_day, report_date = filter_report_day(df_valid, args.date)

    #4. reconcile the counts
    by_segment, by_status = summarise(df_day)

    #5. write the formatted reports + the exceptions file
    write_excel(df_day, by_segment, by_status, report_date, REPORT_XLSX_PATH)
    write_docx(df_day, by_segment, report_date, REPORT_DOCX_PATH)
    save_exceptions(df_exceptions, EXCEPTIONS_PATH)

    print("\ndone.")
